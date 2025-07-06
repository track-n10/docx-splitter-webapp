from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Inches
import tempfile
import os
import zipfile
import re
from io import BytesIO
import json

app = Flask(__name__)
CORS(app)

def clean_text(text):
    """Clean text by removing extra whitespace and normalizing"""
    if not text:
        return ""
    return re.sub(r'\s+', ' ', text.strip())

def find_chapter_breaks(doc, chapter_titles):
    """Find paragraph indices where chapters should be split"""
    chapter_breaks = []
    chapter_titles_lower = [title.lower().strip() for title in chapter_titles]
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = clean_text(paragraph.text).lower()
        
        # Check if this paragraph matches any chapter title
        for j, title in enumerate(chapter_titles_lower):
            if title in text and len(text) <= len(title) * 1.5:  # Allow some flexibility
                chapter_breaks.append({
                    'paragraph_index': i,
                    'chapter_index': j,
                    'title': chapter_titles[j],
                    'found_text': paragraph.text
                })
                break
    
    return sorted(chapter_breaks, key=lambda x: x['paragraph_index'])

def extract_chapter_content(doc, start_idx, end_idx):
    """Extract paragraphs between start and end indices"""
    paragraphs = []
    for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        if para.text.strip():  # Only include non-empty paragraphs
            paragraphs.append({
                'text': para.text,
                'style': para.style.name if para.style else 'Normal'
            })
    return paragraphs

def create_chapter_document(title, paragraphs):
    """Create a new DOCX document for a chapter"""
    new_doc = Document()
    
    # Add title
    title_para = new_doc.add_heading(title, level=1)
    
    # Add content paragraphs
    for para_data in paragraphs:
        para = new_doc.add_paragraph(para_data['text'])
        # Try to preserve some basic styling
        if 'heading' in para_data['style'].lower():
            para.style = 'Heading 2'
    
    return new_doc

@app.route('/')
def index():
    return """
    <h1>DOCX Chapter Splitter API</h1>
    <p>Backend API for processing DOCX files</p>
    <p>Available endpoints:</p>
    <ul>
        <li>POST /api/split - Split DOCX into chapters</li>
        <li>GET /health - Health check</li>
    </ul>
    """

@app.route('/health')
def health():
    return jsonify({'status': 'healthy', 'message': 'API is running'})

@app.route('/api/split', methods=['POST'])
def split_document():
    try:
        # Check if file was uploaded
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        chapter_titles = request.form.get('chapter_titles', '').split('\n')
        chapter_titles = [title.strip() for title in chapter_titles if title.strip()]
        
        if not chapter_titles:
            return jsonify({'error': 'No chapter titles provided'}), 400
        
        # Load the document
        doc = Document(file)
        
        # Find chapter breaks
        chapter_breaks = find_chapter_breaks(doc, chapter_titles)
        
        if not chapter_breaks:
            # If no breaks found, try to split equally
            total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
            para_per_chapter = max(1, total_paragraphs // len(chapter_titles))
            
            chapter_breaks = []
            for i, title in enumerate(chapter_titles):
                chapter_breaks.append({
                    'paragraph_index': i * para_per_chapter,
                    'chapter_index': i,
                    'title': title,
                    'found_text': f'Auto-split at paragraph {i * para_per_chapter}'
                })
        
        # Extract chapters
        chapters = []
        for i, break_point in enumerate(chapter_breaks):
            start_idx = break_point['paragraph_index']
            end_idx = chapter_breaks[i + 1]['paragraph_index'] if i + 1 < len(chapter_breaks) else len(doc.paragraphs)
            
            content = extract_chapter_content(doc, start_idx, end_idx)
            
            chapters.append({
                'title': break_point['title'],
                'content': content,
                'paragraph_count': len(content),
                'found_at': break_point['found_text']
            })
        
        return jsonify({
            'success': True,
            'chapters': chapters,
            'total_chapters': len(chapters),
            'processing_info': {
                'total_paragraphs': len(doc.paragraphs),
                'chapter_breaks_found': len(chapter_breaks)
            }
        })
        
    except Exception as e:
        return jsonify({'error': f'Processing failed: {str(e)}'}), 500

@app.route('/api/download_chapter', methods=['POST'])
def download_chapter():
    try:
        data = request.get_json()
        title = data.get('title', 'Chapter')
        content = data.get('content', [])
        
        # Create new document
        doc = create_chapter_document(title, content)
        
        # Save to memory
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=f'{title}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': f'Download failed: {str(e)}'}), 500

@app.route('/api/download_all', methods=['POST'])
def download_all_chapters():
    try:
        data = request.get_json()
        chapters = data.get('chapters', [])
        
        # Create ZIP file in memory
        zip_buffer = BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for i, chapter in enumerate(chapters):
                # Create document for each chapter
                doc = create_chapter_document(chapter['title'], chapter['content'])
                
                # Save document to memory
                doc_buffer = BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                
                # Add to ZIP
                filename = f"Chapter_{i+1}_{chapter['title'][:30]}.docx"
                # Clean filename
                filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
                zip_file.writestr(filename, doc_buffer.getvalue())
        
        zip_buffer.seek(0)
        
        return send_file(
            zip_buffer,
            as_attachment=True,
            download_name='chapters.zip',
            mimetype='application/zip'
        )
        
    except Exception as e:
        return jsonify({'error': f'ZIP creation failed: {str(e)}'}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)